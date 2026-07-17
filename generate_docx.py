from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_data_dictionary():
    doc = Document()
    doc.add_heading('Online Store Management System (OSMS) - Data Dictionary', 0)

    models = [
        {
            "name": "SignupUser",
            "description": "Core registered users.",
            "fields": [
                ["firstname", "String", "Yes", "First name"],
                ["lastname", "String", "Yes", "Last name"],
                ["mobile", "String", "Yes", "Mobile number"],
                ["email", "String", "Yes (Unique)", "Email address"],
                ["password", "String", "Yes", "Hashed password"],
                ["dob", "String", "No", "Date of birth"],
                ["gender", "String", "No", "Gender"],
                ["role", "String", "user", "Role (user/admin)"],
                ["status", "String", "Active", "Account status"]
            ]
        },
        {
            "name": "LoginUser",
            "description": "Tracks session login events.",
            "fields": [
                ["email", "String", "Yes", "Login email"],
                ["password", "String", "Yes", "Password used"],
                ["loginAt", "Date", "Now", "Login timestamp"]
            ]
        },
        {
            "name": "Product",
            "description": "Store inventory items.",
            "fields": [
                ["title", "String", "Yes", "Product title"],
                ["description", "String", "Yes", "Full description"],
                ["price", "Number", "Yes", "Sale price"],
                ["comparePrice", "Number", "0", "Comparison price"],
                ["stockQuantity", "Number", "0", "Inventory count"],
                ["inStock", "Boolean", "True", "Availability flag"],
                ["image", "String", "Yes", "Main image path"],
                ["category", "String", "Yes", "Category name"],
                ["vendor", "String", "Local", "Vendor name"],
                ["tags", "Array", "No", "Metadata tags"],
                ["sizes", "Array", "No", "Available sizes"],
                ["colors", "Array", "No", "Available colors"],
                ["isSale", "Boolean", "False", "Sale flag"],
                ["images", "Array", "No", "Gallery images"],
                ["productType", "String", "General", "Type classification"],
                ["rating", "Number", "0", "Average stars"],
                ["numReviews", "Number", "0", "Review count"],
                ["viewCount", "Number", "0", "Page views"],
                ["soldCount", "Number", "0", "Total sales"]
            ]
        },
        {
            "name": "Order",
            "description": "Transaction records.",
            "fields": [
                ["user", "ObjectId", "Ref: SignupUser", "Customer ID"],
                ["items", "Array", "Yes", "Items (Product, Title, Price, Qty, Size, Image)"],
                ["shippingAddress", "Object", "Yes", "Full destination address"],
                ["paymentMethod", "String", "Yes", "Payment type"],
                ["paymentResult", "Object", "No", "Gateway results (ID, Status, Time)"],
                ["totalPrice", "Number", "Yes", "Transaction total"],
                ["isPaid", "Boolean", "False", "Payment flag"],
                ["paidAt", "Date", "No", "Payment timestamp"],
                ["isDelivered", "Boolean", "False", "Delivery flag"],
                ["deliveredAt", "Date", "No", "Delivery timestamp"],
                ["orderId", "String", "Yes (Unique)", "Unique Order ID"],
                ["status", "String", "Processing", "Lifecycle status"]
            ]
        },
        {
            "name": "PageContent",
            "description": "CMS Page contents.",
            "fields": [
                ["page", "String", "Yes (Unique)", "Page slug (about/contact/etc)"],
                ["title", "String", "Yes", "Page heading"],
                ["sections", "Mixed", "Yes", "Flexible content blocks"],
                ["lastUpdated", "Date", "Now", "Update timestamp"]
            ]
        },
        {
            "name": "Category",
            "description": "Grouping categories.",
            "fields": [
                ["name", "String", "Yes (Unique)", "Category title"],
                ["image", "String", "Yes", "Icon/Cover image"]
            ]
        },
        {
            "name": "Address",
            "description": "User shipping addresses.",
            "fields": [
                ["user", "ObjectId", "Ref: SignupUser", "Owner ID"],
                ["name", "String", "Yes", "Recipient name"],
                ["mobile", "String", "Yes", "Contact number"],
                ["pincode", "String", "Yes", "Zip code"],
                ["state", "String", "Yes", "State"],
                ["address", "String", "Yes", "Street address"],
                ["locality", "String", "Yes", "Area"],
                ["city", "String", "Yes", "City"],
                ["type", "String", "Home", "Type (HOME/WORK)"],
                ["isDefault", "Boolean", "False", "Default address flag"]
            ]
        },
        {
            "name": "Review",
            "description": "Customer feedback.",
            "fields": [
                ["user", "String", "Yes", "User name/email"],
                ["product", "ObjectId", "Ref: Product", "Product ID"],
                ["rating", "Number", "1-5", "Score"],
                ["title", "String", "No", "Review heading"],
                ["comment", "String", "Yes", "Feedback text"]
            ]
        },
        {
            "name": "Blog",
            "description": "Marketing articles.",
            "fields": [
                ["title", "String", "Yes", "Heading"],
                ["category", "String", "Yes", "Blog category"],
                ["author", "String", "Yes", "Writer"],
                ["date", "String", "Yes", "Post date"],
                ["excerpt", "String", "Yes", "Brief summary"],
                ["content", "String", "Yes", "Main Body"],
                ["image", "String", "Yes", "Graphic URL"],
                ["tags", "Array", "No", "Keywords"]
            ]
        },
        {
            "name": "Contact",
            "description": "Inbound queries.",
            "fields": [
                ["name", "String", "Yes", "Sender name"],
                ["email", "String", "Yes", "Contact email"],
                ["phone", "String", "Yes", "Contact phone"],
                ["subject", "String", "Yes", "Topic"],
                ["message", "String", "Yes", "Content body"]
            ]
        }
    ]

    for model in models:
        doc.add_heading(f"Collection: {model['name']}", level=1)
        doc.add_paragraph(model['description'])
        
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Field Name'
        hdr_cells[1].text = 'Data Type'
        hdr_cells[2].text = 'Mandatory/Constraint'
        hdr_cells[3].text = 'Description'

        # Make header bold
        for cell in hdr_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True

        for field in model['fields']:
            row_cells = table.add_row().cells
            row_cells[0].text = field[0]
            row_cells[1].text = field[1]
            row_cells[2].text = field[2]
            row_cells[3].text = field[3]
        
        doc.add_paragraph("\n")

    doc.save('OSMS_Data_Dictionary_v2.docx')
    print("OSMS_Data_Dictionary_v2.docx created with Table formatting.")

def create_database_tables():
    doc = Document()
    doc.add_heading('Online Store Management System (OSMS) - Database Tables Structure', 0)

    tables = [
        ("SignupUsers", ["_id (PK)", "firstname", "lastname", "email (UQ)", "mobile", "password", "role", "status", "createdAt", "updatedAt"]),
        ("LoginUsers", ["_id (PK)", "email", "password", "loginAt", "createdAt"]),
        ("Products", ["_id (PK)", "title", "description", "price", "stockQuantity", "category", "image", "vendor", "isSale", "rating", "numReviews"]),
        ("Orders", ["_id (PK)", "user (FK)", "orderId (UQ)", "items[]", "totalPrice", "status", "paymentMethod", "isPaid", "paidAt", "createdAt"]),
        ("PageContents", ["_id (PK)", "page (UQ)", "title", "sections", "lastUpdated", "createdAt"]),
        ("Categories", ["_id (PK)", "name (UQ)", "image", "createdAt"]),
        ("Addresses", ["_id (PK)", "user (FK)", "name", "mobile", "pincode", "state", "address", "city", "type", "isDefault"]),
        ("Reviews", ["_id (PK)", "user", "product (FK)", "rating", "title", "comment", "createdAt"]),
        ("Blogs", ["_id (PK)", "title", "category", "author", "date", "image", "content", "tags[]"]),
        ("Contacts", ["_id (PK)", "name", "email", "phone", "subject", "message", "createdAt"])
    ]

    for table_name, columns in tables:
        doc.add_heading(f"Table: {table_name}", level=2)
        
        # Creating a proper table for column list
        table = doc.add_table(rows=1, cols=1)
        table.style = 'Table Grid'
        table.rows[0].cells[0].text = "Columns / Fields"
        table.rows[0].cells[0].paragraphs[0].runs[0].bold = True
        
        row_cells = table.add_row().cells
        row_cells[0].text = ", ".join(columns)
        
        doc.add_paragraph("\n")

    doc.save('OSMS_Database_Tables_v2.docx')
    print("OSMS_Database_Tables_v2.docx updated with missing tables.")

if __name__ == "__main__":
    create_data_dictionary()
    create_database_tables()
